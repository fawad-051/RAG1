# app4.py  -- Advanced RAG Q&A (upgraded)
# Requirements (suggested):
# pip install streamlit langchain-groq langchain-core langchain-community langchain-text-splitters langchain-chroma sentence-transformers==2.2.2 scikit-learn python-docx python-dotenv

import os
import tempfile
import streamlit as st
import time
import shutil
import atexit
import uuid
import json
from datetime import datetime
from dotenv import load_dotenv

# LangChain / embedding / vectorstore imports
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Try different embedding approaches
try:
    # First try: use HuggingFaceEmbeddings with a simpler model
    from langchain_community.embeddings import HuggingFaceEmbeddings
    EMBEDDINGS_AVAILABLE = True
except Exception as e:
    st.warning(f"HuggingFaceEmbeddings not available: {e}")
    EMBEDDINGS_AVAILABLE = False

try:
    from langchain_chroma import Chroma
    CHROMA_AVAILABLE = True
except Exception as e:
    st.warning(f"Chroma not available: {e}")
    CHROMA_AVAILABLE = False

from langchain_core.output_parsers import StrOutputParser

# Optional imports
try:
    from sklearn.cluster import KMeans
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Load environment variables
load_dotenv()

# -----------------------------
# Streamlit basic config / UI
# -----------------------------
st.set_page_config(page_title="Advanced RAG Q&A", page_icon="üìö", layout="wide")
st.title("üöÄ Advanced RAG Q&A ‚Äî Upgraded")

# Debug toggle
DEBUG = st.sidebar.checkbox("Debug mode (console logs)", False)

def debug_log(msg):
    if DEBUG:
        print("DEBUG:", msg)

# -----------------------------
# Session & storage helpers
# -----------------------------
DEFAULT_BASE_DIR = "./rag_data"
os.makedirs(DEFAULT_BASE_DIR, exist_ok=True)

def session_dir(session_id):
    d = os.path.join(DEFAULT_BASE_DIR, f"session_{session_id}")
    os.makedirs(d, exist_ok=True)
    return d

def chat_history_path(session_id):
    return os.path.join(session_dir(session_id), "chat_history.json")

def save_chat_history(session_id, messages):
    try:
        with open(chat_history_path(session_id), "w", encoding="utf-8") as f:
            json.dump({"messages": messages, "updated_at": datetime.utcnow().isoformat()}, f, ensure_ascii=False, indent=2)
        debug_log(f"Saved chat history for {session_id}")
    except Exception as e:
        debug_log(f"Error saving chat history: {e}")

def load_chat_history(session_id):
    p = chat_history_path(session_id)
    if os.path.exists(p):
        try:
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("messages", [])
        except Exception as e:
            debug_log(f"Error loading chat history: {e}")
    return []

# -----------------------------
# Sidebar configuration
# -----------------------------
with st.sidebar:
    st.header("‚öôÔ∏è Configuration")
    session_id = st.text_input("Session ID", value=os.getenv("DEFAULT_SESSION_ID", "default_session"))
    api_key_input = st.text_input("Groq API Key", type="password")
    model_choice = st.selectbox("Model", ["llama-3.3-70b-versatile", "llama-3.3-70b", "openai/gpt-oss-120b"], index=0)
    top_k = st.slider("Top K Chunks", 1, 12, 4)
    similarity_threshold = st.slider("Similarity Threshold (distance)", 0.0, 1.0, 0.7)
    debug_checkbox = st.checkbox("Show more debug info", False)
    
    # Embedding method selection
    embedding_method = st.selectbox(
        "Embedding Method",
        ["HuggingFace (Local)", "OpenAI (API)"],
        index=0,
        help="HuggingFace uses local models, OpenAI requires API key"
    )
    
    if embedding_method == "OpenAI (API)":
        openai_api_key = st.text_input("OpenAI API Key", type="password")
    else:
        openai_api_key = None
    
    if st.button("üßπ Cleanup Vector Store (this session)"):
        # remove vectorstore for session
        idx_dir = os.path.join(session_dir(session_id), "chroma_index")
        if os.path.exists(idx_dir):
            shutil.rmtree(idx_dir, ignore_errors=True)
            st.success("Cleaned session vector store.")
        else:
            st.info("No vectorstore found for this session.")

# link debug flags
if debug_checkbox:
    DEBUG = True

# Use API key provided or from .env
api_key = api_key_input or os.getenv("GROQ_API_KEY")
if not api_key:
    st.warning("‚ö†Ô∏è Please enter your Groq API key in the sidebar or in a .env file.")
    st.stop()

# Initialize LLM
try:
    llm = ChatGroq(
        groq_api_key=api_key,
        model_name=model_choice,
        temperature=0.1
    )
except Exception as e:
    st.error(f"Failed to initialize Groq model: {e}")
    st.stop()

# -----------------------------
# Embeddings setup with fallback
# -----------------------------
@st.cache_resource(show_spinner=False)
def get_embeddings(_method="HuggingFace (Local)", _openai_key=None):
    if _method == "OpenAI (API)" and _openai_key:
        try:
            from langchain_openai import OpenAIEmbeddings
            return OpenAIEmbeddings(openai_api_key=_openai_key)
        except Exception as e:
            st.warning(f"OpenAI embeddings failed: {e}. Falling back to HuggingFace.")
    
    # Fallback to HuggingFace
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
        # Use a simpler, more compatible model
        return HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},  # Force CPU for compatibility
            encode_kwargs={'normalize_embeddings': False}
        )
    except Exception as e:
        st.error(f"Failed to initialize embeddings: {e}")
        # Ultimate fallback - use a dummy embeddings (not recommended for production)
        st.warning("Using dummy embeddings as last resort. This is not suitable for production.")
        from langchain_community.embeddings import FakeEmbeddings
        return FakeEmbeddings(size=384)

try:
    embeddings = get_embeddings(embedding_method, openai_api_key)
except Exception as e:
    st.error(f"Embeddings initialization error: {e}")
    st.stop()

# -----------------------------
# File upload (PDF, txt, docx)
# -----------------------------
st.header("üì§ Upload documents (PDF / TXT / DOCX)")
uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "docx"], accept_multiple_files=True)
if not uploaded_files:
    st.info("Upload one or more files to begin. You can also drag multiple files.")
    # Load previous chat history and show quick actions if available
    previous_msgs = load_chat_history(session_id)
    if previous_msgs:
        st.success(f"Loaded saved chat for session '{session_id}' ({len(previous_msgs)} messages).")
        if st.button("üîÅ Load chat into UI"):
            st.session_state.messages = previous_msgs
    st.stop()

# -----------------------------
# Temporary save uploaded files and load docs
# -----------------------------
# Document loader: prefer PyMuPDFLoader; fallback to PyPDFLoader
try:
    from langchain_community.document_loaders import PyMuPDFLoader as PDFLoader
except Exception:
    try:
        from langchain_community.document_loaders import PyPDFLoader as PDFLoader
    except Exception:
        PDFLoader = None

from langchain_core.documents import Document

all_docs = []
tmp_paths = []
with st.spinner("Processing uploaded files..."):
    for f in uploaded_files:
        name = f.name
        suffix = os.path.splitext(name)[1].lower()
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(f.getvalue())
            tmp_paths.append(tmp.name)
            tmp_name = tmp.name

        try:
            if suffix == ".pdf":
                if PDFLoader is None:
                    st.error("PDF processing is not available. Please install PyMuPDF or PyPDF2.")
                    continue
                loader = PDFLoader(tmp_name)
                docs = loader.load()
            elif suffix == ".txt":
                text = open(tmp_name, "r", encoding="utf-8", errors="ignore").read()
                docs = [Document(page_content=text, metadata={"source_file": name})]
            elif suffix == ".docx" and DOCX_AVAILABLE:
                doc = docx.Document(tmp_name)
                full = "\n".join([p.text for p in doc.paragraphs])
                docs = [Document(page_content=full, metadata={"source_file": name})]
            else:
                # fallback: plain read
                text = open(tmp_name, "r", encoding="utf-8", errors="ignore").read()
                docs = [Document(page_content=text, metadata={"source_file": name})]

            # ensure metadata tags
            for d in docs:
                d.metadata["source_file"] = name
            all_docs.extend(docs)
        except Exception as e:
            st.error(f"Failed to load {name}: {e}")

# cleanup temp files
for p in tmp_paths:
    try:
        os.unlink(p)
    except Exception:
        pass

if not all_docs:
    st.error("No readable text found in uploaded documents. Try different files.")
    st.stop()

st.success(f"‚úÖ Loaded {len(all_docs)} pages/chunks from {len(uploaded_files)} file(s)")

# -----------------------------
# Split into chunks
# -----------------------------
splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=120)
splits = splitter.split_documents(all_docs)
st.info(f"üìÑ Created {len(splits)} text chunks.")

# -----------------------------
# Vectorstore (persist per-session)
# -----------------------------
INDEX_DIR = os.path.join(session_dir(session_id), "chroma_index")

@st.cache_resource(show_spinner=False)
def init_vectorstore(_splits, _embeddings, persist_dir):
    if not CHROMA_AVAILABLE:
        st.error("Chroma vectorstore is not available. Please check dependencies.")
        return None
        
    # remove previous chroma in this session to avoid stale corruption
    try:
        vs = Chroma.from_documents(_splits, _embeddings, persist_directory=persist_dir)
        return vs
    except Exception as e:
        # attempt to cleanup then recreate
        try:
            if os.path.exists(persist_dir):
                shutil.rmtree(persist_dir, ignore_errors=True)
            vs = Chroma.from_documents(_splits, _embeddings, persist_directory=persist_dir)
            return vs
        except Exception as e2:
            st.error(f"Failed to init Chroma vectorstore: {e2}")
            return None

with st.spinner("Initializing vector store..."):
    try:
        vectorstore = init_vectorstore(splits, embeddings, INDEX_DIR)
        if vectorstore is None:
            st.error("Failed to initialize vector store. Check the logs for details.")
            st.stop()
    except Exception as e:
        st.error(f"Vectorstore initialization error: {e}")
        st.stop()

st.session_state.vectorstore_initialized = True

# -----------------------------
# Retriever with threshold filtering
# -----------------------------
def create_retriever(vs, k=5, distance_threshold=0.7):
    def retrieve(question):
        try:
            docs_with_scores = vs.similarity_search_with_score(question, k=k*3)
            # docs_with_scores -> list of (doc, distance)
            filtered = [d for d, dist in docs_with_scores if dist <= distance_threshold]
            if not filtered:
                filtered = [d for d, _ in docs_with_scores[:k]]
            return filtered[:k]
        except Exception as e:
            debug_log(f"Retriever internal error: {e}")
            return vs.similarity_search(question, k=k)
    return retrieve

retriever = create_retriever(vectorstore, k=top_k, distance_threshold=similarity_threshold)

# -----------------------------
# RAG chain definition
# -----------------------------
def rag_chain(question, chat_history_messages):
    start_time = time.time()
    docs = retriever(question)
    if not docs:
        return "‚ö†Ô∏è No relevant text found in your documents.", [], 0.0

    context = "\n\n".join([d.page_content for d in docs])
    qa_prompt = ChatPromptTemplate.from_messages([
        ("system",
         "You are a helpful assistant. Answer based **only** on the provided context.\n"
         "If the answer is not in the context, say: 'The answer is not available in the provided documents.'\n\n"
         "Context:\n{context}\n"),
        ("human", "Question: {question}")
    ])

    # chain style: create prompt, call llm
    chain = qa_prompt | llm | StrOutputParser()
    try:
        response = chain.invoke({"context": context, "question": question})
    except Exception as e:
        response = f"Error generating response: {e}"
    duration = round(time.time() - start_time, 2)
    return response, docs, duration

# -----------------------------
# Insights & clustering
# -----------------------------
def generate_document_summary(llm, docs, max_chars=4000):
    # join some content and ask model for short summary
    preview = "\n\n".join([d.page_content[:1000] for d in docs[:10]])
    if len(preview) > max_chars:
        preview = preview[:max_chars]
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a helpful assistant summarizer."),
        ("human", f"Summarize the following documents into a concise set of bullet points (max 6 bullets). Text:\n\n{preview}")
    ])
    chain = prompt | llm | StrOutputParser()
    try:
        return chain.invoke({})
    except Exception as e:
        return f"Unable to generate summary: {e}"

def cluster_chunks(_embeddings, chunks, n_clusters=4):
    if not SKLEARN_AVAILABLE:
        return None
    try:
        # Use a simpler approach to avoid complex dependencies
        embs = [_embeddings.embed_query(c.page_content) for c in chunks[:50]]  # Limit for performance
        import numpy as np
        em = np.array(embs)
        kmeans = KMeans(n_clusters=min(n_clusters, len(em)), random_state=0).fit(em)
        clusters = {}
        for idx, label in enumerate(kmeans.labels_):
            clusters.setdefault(int(label), []).append(chunks[idx])
        return clusters
    except Exception as e:
        debug_log(f"Clustering failed: {e}")
        return None

# -----------------------------
# Chat UI & interaction
# -----------------------------
if "messages" not in st.session_state:
    # load saved chat history by session id if present
    previous = load_chat_history(session_id)
    if previous:
        st.session_state.messages = previous
    else:
        st.session_state.messages = []

if "chathistory" not in st.session_state:
    st.session_state.chathistory = {}

def get_chat_history_obj(sid):
    if sid not in st.session_state.chathistory:
        st.session_state.chathistory[sid] = ChatMessageHistory()
    return st.session_state.chathistory[sid]

# Display existing messages
st.header("üí¨ Chat with your documents")
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        st.markdown(m["content"])

# Show insights block (collapsible)
with st.expander("üîé Document Insights"):
    st.write(f"**Files uploaded:** {', '.join({d.metadata.get('source_file','?') for d in all_docs})}")
    st.write(f"**Total raw pages/chunks:** {len(all_docs)}")
    st.write(f"**Chunks created:** {len(splits)}")
    if st.button("üßæ Generate quick summary (using model)"):
        with st.spinner("Generating summary..."):
            summary = generate_document_summary(llm, all_docs)
            st.markdown(summary)

    if SKLEARN_AVAILABLE:
        if st.button("üìÇ Try topic clustering of chunks"):
            with st.spinner("Clustering chunks..."):
                clusters = cluster_chunks(embeddings, splits, n_clusters=4)
                if clusters:
                    for k, items in clusters.items():
                        st.markdown(f"**Cluster {k} ‚Äî {len(items)} chunks**")
                        st.write(items[0].page_content[:300] + "...")
                else:
                    st.info("Clustering did not produce results.")
    else:
        st.info("Topic clustering requires scikit-learn (`pip install scikit-learn`).")

# Chat input
user_q = st.chat_input("Ask something about your uploaded documents...")

if user_q:
    # Add user message to UI & history
    history_obj = get_chat_history_obj(session_id)
    with st.chat_message("user"):
        st.markdown(user_q)
    st.session_state.messages.append({"role": "user", "content": user_q})
    history_obj.add_user_message(user_q)

    # Generate response
    with st.chat_message("assistant"):
        with st.spinner("Analyzing documents and generating answer..."):
            try:
                response, docs, duration = rag_chain(user_q, history_obj)
                st.markdown(response)
                st.caption(f"‚è±Ô∏è Answer generated in {duration} seconds")

                # Source display
                if docs:
                    with st.expander("üìÇ Sources Used (click to expand)"):
                        for d in docs:
                            src = d.metadata.get("source_file", "Unknown File")
                            st.markdown(f"**üìÑ {src}** ‚Äî Preview:")
                            # highlight small context: we simply show snippet with user's query bolded
                            snippet = d.page_content[:800]
                            # naive highlight: replace occurrences (case-insensitive)
                            try:
                                import re
                                pattern = re.compile(re.escape(user_q[:60]), re.IGNORECASE)
                                snippet = pattern.sub(lambda m: f"**{m.group(0)}**", snippet)
                            except Exception:
                                pass
                            st.write(snippet + "...")
                            st.divider()

                # Save assistant message
                st.session_state.messages.append({"role": "assistant", "content": response})
                history_obj.add_ai_message(response)
                # persist chat to file
                save_chat_history(session_id, st.session_state.messages)

            except Exception as e:
                err = f"Error while generating response: {e}"
                st.error(err)
                st.session_state.messages.append({"role": "assistant", "content": err})
                save_chat_history(session_id, st.session_state.messages)

# -----------------------------
# Chat export / download
# -----------------------------
st.header("üì• Export / Session Tools")
col1, col2, col3 = st.columns(3)

with col1:
    if st.button("üì• Download Chat as TXT"):
        chat_text = "\n".join([f"{m['role'].upper()}: {m['content']}" for m in st.session_state.messages])
        st.download_button("Download TXT", data=chat_text, file_name=f"chat_{session_id}.txt", mime="text/plain")

with col2:
    if st.button("üíæ Export Chat as JSON"):
        chat_json = json.dumps(st.session_state.messages, ensure_ascii=False, indent=2)
        st.download_button("Download JSON", data=chat_json, file_name=f"chat_{session_id}.json", mime="application/json")

with col3:
    if st.button("üóëÔ∏è Clear current chat (session)"):
        st.session_state.messages = []
        # remove persisted file
        p = chat_history_path(session_id)
        if os.path.exists(p):
            os.remove(p)
        st.rerun()

# -----------------------------
# Admin & diagnostics
# -----------------------------
with st.expander("üìä Admin / Diagnostics"):
    st.write(f"Vectorstore directory: {INDEX_DIR}")
    try:
        idx_exists = os.path.exists(INDEX_DIR)
        st.write(f"Vectorstore exists: {idx_exists}")
    except Exception:
        pass
    st.write(f"Documents loaded: {len(all_docs)}")
    st.write(f"Chunks: {len(splits)}")
    st.write(f"Session ID: {session_id}")
    st.write(f"Embedding method: {embedding_method}")

# -----------------------------
# Atexit cleanup safety (optional)
# -----------------------------
def cleanup_vectorstore_on_exit():
    # Do NOT auto-delete ‚Äî we persist per session. This function kept for manual use.
    debug_log("Exiting app. If you want to cleanup vectorstores, use the sidebar button.")

atexit.register(cleanup_vectorstore_on_exit)
